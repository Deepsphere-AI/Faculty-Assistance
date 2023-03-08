import streamlit as st
from streamlit_option_menu import option_menu
import os
import io
import docx 
import source.title_1 as head
import QuestionGeneration.text_file as txt
import QuestionGeneration.split_prompt as spilt
import source.GPT_3 as gpt
import QuestionGeneration.pdf as pdf
import QuestionGeneration.pdf_textretrive as pdf_retrive
import openai
from docx import document
import QuestionGeneration.word_doc_que_gen as word_que
def prev():
    st.session_state['preview']="No"
    st.session_state['preview2']="No"
def quens_gen():
    head.title()
    st.markdown("<p style='text-align: center; color: black; font-size:20px;'><span style='font-weight: bold'>Problem Statement: </span>Application to Question Generator    </p>", unsafe_allow_html=True)
    st.markdown("<hr style=height:2.5px;background-color:gray>",unsafe_allow_html=True)
    w1,col1,col2,w2=st.columns((1,1.5,2.5,1))
    w12,col11,col22,w22=st.columns((1,1.5,2.5,1))
    c2,c1,c3=st.columns((1.5,3,1))
    bc1,bc2,bc3=st.columns((10,5,10))
    cc2,cc1,cc3=st.columns((1,4,1))
    openai.api_key = os.environ["API_KEY"]
    with col1:
        st.markdown("# ")
        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Input Type</span></p>", unsafe_allow_html=True)
        #st.write("# Input Typeㅤㅤㅤ")   
    with col2:
        vAR_input_file_select = st.selectbox("",["Select","Files","Passage"],key="Clear_quetype")
        # file operations
        if  vAR_input_file_select == "Files":
            with col1:
                st.write("# ")
                st.markdown("## ")
                st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Upload Here</span></p>", unsafe_allow_html=True)
            with col2:
                prompt_file_txt = st.file_uploader("", type=['pdf', 'txt'])
            if 'txt' in str(prompt_file_txt):
                vAR_content=prompt_file_txt.getvalue().decode("utf-8")
                with col11:
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Preview</span></p>", unsafe_allow_html=True)
                with col22:
                    vAR_preview = st.selectbox("",["Select","Yes","No"],key="preview")
                if vAR_preview =="Yes":
                    with cc1:
                        st.write(vAR_content)
                with col11:
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Question Type</span></p>", unsafe_allow_html=True)
                with col22:
                    vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile",on_change=prev)   
                if prompt_file_txt:
                    txt_content = prompt_file_txt.getvalue().decode("utf-8")
                    to_split_content=spilt.fullstr(txt_content)
                    os.remove("Result//newfile.txt")
                    if vAR_input_quetype !="Select":
                        with col22:    
                            st.markdown("")
                            button_placeholder = st.empty()
                            button_clicked = False
                            key=0
                            while not button_clicked:
                                key=key+1
                                button_clicked = button_placeholder.button('Submit',key=key)
                                break
                            if button_clicked:
                                button_placeholder.empty()
                                #st.button("Submit"):
                                #Creating a new word document
                                vAR_new_doc = docx.Document()
                                fin_txt = ""
                                with col22:
                                    st.info("Extract text from file")
                                    st.info("Generating "+vAR_input_quetype)
                                no_iter=len(to_split_content)
                                half=no_iter//2
                                vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                serial=0
                                for split_content in to_split_content:
                                    if split_content==to_split_content[half]:
                                        with col22:
                                            st.info("Half of the process is completed")
                                    serial_no = vAR_x[serial]
                                    serial+=1
                                    txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                    vAR_responce = gpt.generate_response3(txt_prompt)
                                    st.write(vAR_responce)
                                    fin_txt =fin_txt + vAR_responce
                                with col22:
                                    st.info("Response are storing in word file")
                                vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                #vAR_final_docs.save(str(vAR_input_quetype)+".docx")
                                bio = io.BytesIO()
                                vAR_final_docs.save(bio)
                                # with bc2:
                                #     st.markdown("")
                                #     st.button("Clear", on_click=cr.clear_text)
                                with col22:
                                    st.markdown("")
                                    st.download_button(
                                        label="Download",
                                        data=bio.getvalue(),
                                        file_name=vAR_input_quetype+".docx",
                                        mime="docx"
                                    )
                                
################################################################################################################            
            # pdf file operatons 
            if 'pdf' in str(prompt_file_txt):
                pdf.header_footer_cuter(prompt_file_txt)
                pdf_content = pdf_retrive.pdf_text()
                with col11:
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Preview</span></p>", unsafe_allow_html=True)
                with col22:
                    vAR_preview = st.selectbox("",["Select","Yes","No"],key="preview2")
                if vAR_preview =="Yes":
                    with cc1:
                        # pdf.header_footer_cuter(prompt_file_txt)
                        # pdf_content = pdf_retrive.pdf_text()
                        st.write(pdf_content)
                with col11:
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Question Type</span></p>", unsafe_allow_html=True)
                with col22:
                    vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile",on_change=prev) 
                if vAR_input_quetype !="Select":
                    with col22:
                        st.markdown("")
                        button_placeholder = st.empty()
                        button_clicked = False
                        key=0
                        while not button_clicked:
                            key=key+1
                            button_clicked = button_placeholder.button('Submit',key=key)
                            break
                        if button_clicked:
                            button_placeholder.empty()
                            with col22:
                                st.info("Removing the header and footer of the pdf")
                                
                            if prompt_file_txt is not None:
                                with col22:
                                    st.info("Extract text from file")
                                to_split_content=spilt.fullstr(pdf_content)
                                os.remove("Result//newfile.txt")
                                with col22:
                                    st.info("Summarizing the content and passes to generate questions")
                                #Creating a new word document.
                                vAR_new_doc = docx.Document()
                                fin_txt = ""
                                no_iter=len(to_split_content)
                                half=no_iter//2
                                with col22:
                                    st.info("Generating "+vAR_input_quetype)
                                vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                serial=0
                                for split_content in to_split_content:
                                    if split_content==to_split_content[half]:
                                        with col22:
                                            st.info("Half of the process is completed")
                                    split_content = gpt.summarizing(split_content)
                                    serial_no = vAR_x[serial]
                                    serial+=1
                                    txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                    vAR_responce = gpt.generate_response3(txt_prompt)
                                    fin_txt =fin_txt + vAR_responce
                                with col22:
                                    st.info("Response are storing in word file")
                                vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                with col22:
                                    st.markdown("")
                                    bio = io.BytesIO()
                                    vAR_final_docs.save(bio)
                                    st.download_button(
                                        label="Download",
                                        data=bio.getvalue(),
                                        file_name=vAR_input_quetype+".docx",
                                        mime="docx"
                                    )
 ##################################################################################################################                              
        # Passage operations
        elif  vAR_input_file_select == "Passage":
            with col1:
                st.markdown("# ")
                st.markdown("# ")
                st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Question Type</span></p>", unsafe_allow_html=True)
            with col1:
                st.markdown("## ")
                st.markdown("# ")
                st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Enter the Passage</span></p>", unsafe_allow_html=True)
            vAR_passage = st.text_area("",key = "Clear_input_para")
            to_split_content=spilt.fullstr(vAR_passage)
            with col2:
                    vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile") 
            if vAR_input_quetype !="Select":
                with col22:    
                    st.markdown("# ")
                    button_placeholder = st.empty()
                    button_clicked = False
                    key=0
                    while not button_clicked:
                        key=key+1
                        button_clicked = button_placeholder.button('Submit',key=key)
                        break
                    if button_clicked:
                        if vAR_passage !="":
                            if len(vAR_passage) >= 500:
                                button_placeholder.empty()
                                #Creating a new word document.
                                vAR_new_doc = docx.Document()
                                fin_txt = ""
                                no_iter=len(to_split_content)
                                half=no_iter//2
                                with col22:
                                    st.info("Generating "+vAR_input_quetype)
                                vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                serial=0
                                for split_content in to_split_content:
                                    if split_content==to_split_content[half]:
                                        with col22:
                                            st.info("Half of the process is completed")
                                    serial_no = vAR_x[serial]
                                    serial += 1
                                    txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                    vAR_responce = gpt.generate_response3(txt_prompt)
                                    fin_txt = fin_txt + vAR_responce
                                vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                with col22:
                                    st.write("# ")
                                    st.info("Response are storing in word file")
                                    #st.success("Response saved to "+vAR_input_quetype+".docx")
                                    #vAR_final_docs.save(str(vAR_input_quetype)+".docx")
                                    bio = io.BytesIO()
                                    vAR_final_docs.save(bio)
                                with col22:
                                    st.markdown("")
                                    st.download_button(
                                        label="Download",
                                        data=bio.getvalue(),
                                        file_name=vAR_input_quetype+".docx",
                                        mime="docx"
                                    )
                            else:
                                with col2:
                                    st.warning("The Content is too Short")
                        else:
                            with col2:
                                st.warning("Enter the Passage")
